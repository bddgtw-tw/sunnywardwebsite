import sys
import os
from docx import Document

def print_docx_content(filepath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
    
    doc = Document(filepath)
    print("=== DOCX CONTENT START ===")
    for para in doc.paragraphs:
        if para.text.strip():
            try:
                print(para.text)
            except UnicodeEncodeError:
                print(para.text.encode('utf-8', errors='ignore').decode('cp950', errors='ignore'))
    
    for table in doc.tables:
        print("\n--- Table ---")
        for row in table.rows:
            text = []
            for cell in row.cells:
                txt = cell.text.strip().replace('\n', ' ')
                text.append(txt)
            line = " | ".join(text)
            try:
                print(line)
            except UnicodeEncodeError:
                print(line.encode('utf-8', errors='ignore').decode('cp950', errors='ignore'))
    print("=== DOCX CONTENT END ===")

if __name__ == "__main__":
    target = r"c:\Users\bddgt\.gemini\antigravity\scratch\sunnyward-website\data\2026.7.16_website_content_editing.docx"
    print_docx_content(target)
